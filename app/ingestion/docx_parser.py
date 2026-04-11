from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

@dataclass(slots=True)
class ParagraphElement:
    element_id: str
    element_type: str
    section_path: list[str]
    style: str
    text: str
    formula_image_ids: list[str] = field(default_factory=list)


@dataclass(slots=True)
class TableCellElement:
    row_index: int
    col_index: int
    text: str
    formula_image_ids: list[str] = field(default_factory=list)


@dataclass(slots=True)
class TableRowElement:
    row_index: int
    values: list[str]
    cells: list[TableCellElement] = field(default_factory=list)
    formula_image_ids: list[str] = field(default_factory=list)


@dataclass(slots=True)
class TableElement:
    element_id: str
    element_type: str
    section_path: list[str]
    row_count: int
    col_count: int
    rows: list[TableRowElement] = field(default_factory=list)


@dataclass(slots=True)
class FormulaImageAsset:
    asset_id: str
    filename: str
    content: bytes = field(repr=False)


@dataclass(slots=True)
class ParsedDocx:
    source_name: str
    paragraph_count: int
    table_count: int
    paragraphs: list[ParagraphElement] = field(default_factory=list)
    tables: list[TableElement] = field(default_factory=list)
    formula_images: list[FormulaImageAsset] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_name": self.source_name,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "paragraphs": [
                {
                    "element_id": item.element_id,
                    "element_type": item.element_type,
                    "section_path": item.section_path,
                    "style": item.style,
                    "text": item.text,
                    "formula_image_ids": item.formula_image_ids,
                }
                for item in self.paragraphs
            ],
            "tables": [
                {
                    "element_id": table.element_id,
                    "element_type": table.element_type,
                    "section_path": table.section_path,
                    "row_count": table.row_count,
                    "col_count": table.col_count,
                    "rows": [
                        {
                            "row_index": row.row_index,
                            "values": row.values,
                            "formula_image_ids": row.formula_image_ids,
                            "cells": [
                                {
                                    "row_index": cell.row_index,
                                    "col_index": cell.col_index,
                                    "text": cell.text,
                                    "formula_image_ids": cell.formula_image_ids,
                                }
                                for cell in row.cells
                            ],
                        }
                        for row in table.rows
                    ],
                }
                for table in self.tables
            ],
            "formula_images": [
                {
                    "asset_id": asset.asset_id,
                    "filename": asset.filename,
                }
                for asset in self.formula_images
            ],
        }


def parse_docx_bytes(source_name: str, content: bytes) -> ParsedDocx:
    document = Document(BytesIO(content))
    section_path: list[str] = []
    paragraphs: list[ParagraphElement] = []
    tables: list[TableElement] = []
    formula_images: list[FormulaImageAsset] = []
    paragraph_index = 0
    table_index = 0
    formula_image_index = 0

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            content_block, formula_image_index, new_assets = _normalize_paragraph_text(block, formula_image_index)
            formula_images.extend(new_assets)
            if not content_block.text:
                continue

            style_name = block.style.name if block.style is not None else "Normal"
            heading_level = _detect_heading_level(style_name, content_block.text, paragraph_index)
            if heading_level is not None:
                section_path = section_path[: max(heading_level - 1, 0)]
                section_path.append(content_block.text)

            paragraph_index += 1
            paragraphs.append(
                ParagraphElement(
                    element_id=f"p-{paragraph_index}",
                    element_type="paragraph",
                    section_path=section_path[:],
                    style=style_name,
                    text=content_block.text,
                    formula_image_ids=content_block.formula_image_ids,
                )
            )
            continue

        if isinstance(block, Table):
            table_index += 1
            table_rows: list[TableRowElement] = []
            col_count = 0

            for row_idx, row in enumerate(block.rows):
                values: list[str] = []
                row_cells: list[TableCellElement] = []
                row_formula_image_ids: list[str] = []

                for col_idx, cell in enumerate(row.cells):
                    cell_content, formula_image_index, new_assets = _normalize_cell_text(cell, formula_image_index)
                    formula_images.extend(new_assets)
                    values.append(cell_content.text)
                    row_formula_image_ids.extend(cell_content.formula_image_ids)
                    row_cells.append(
                        TableCellElement(
                            row_index=row_idx,
                            col_index=col_idx,
                            text=cell_content.text,
                            formula_image_ids=cell_content.formula_image_ids,
                        )
                    )

                col_count = max(col_count, len(values))
                table_rows.append(
                    TableRowElement(
                        row_index=row_idx,
                        values=values,
                        cells=row_cells,
                        formula_image_ids=list(dict.fromkeys(row_formula_image_ids)),
                    )
                )

            tables.append(
                TableElement(
                    element_id=f"t-{table_index}",
                    element_type="table",
                    section_path=section_path[:],
                    row_count=len(table_rows),
                    col_count=col_count,
                    rows=table_rows,
                )
            )

    return ParsedDocx(
        source_name=source_name,
        paragraph_count=len(paragraphs),
        table_count=len(tables),
        paragraphs=paragraphs,
        tables=tables,
        formula_images=formula_images,
    )


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


@dataclass(slots=True)
class NormalizedContent:
    text: str
    formula_image_ids: list[str] = field(default_factory=list)


def _normalize_cell_text(cell: _Cell, formula_image_index: int) -> tuple[NormalizedContent, int, list[FormulaImageAsset]]:
    texts: list[str] = []
    formula_image_ids: list[str] = []
    assets: list[FormulaImageAsset] = []
    current_index = formula_image_index

    for paragraph in cell.paragraphs:
        content, current_index, new_assets = _normalize_paragraph_text(paragraph, current_index)
        if content.text:
            texts.append(content.text)
        formula_image_ids.extend(content.formula_image_ids)
        assets.extend(new_assets)

    return (
        NormalizedContent(text="\n".join(texts), formula_image_ids=list(dict.fromkeys(formula_image_ids))),
        current_index,
        assets,
    )


def _normalize_paragraph_text(paragraph: Paragraph, formula_image_index: int) -> tuple[NormalizedContent, int, list[FormulaImageAsset]]:
    parts = [paragraph.text.strip()]
    parts.extend(_extract_omml_formulas(paragraph))
    image_markers, formula_image_ids, next_index, assets = _extract_formula_image_markers(paragraph, formula_image_index)
    parts.extend(image_markers)
    return NormalizedContent(
        text="\n".join(part for part in parts if part).strip(),
        formula_image_ids=formula_image_ids,
    ), next_index, assets


def _extract_omml_formulas(paragraph: Paragraph) -> list[str]:
    formulas: list[str] = []
    for formula in paragraph._p.xpath(
        './/*[local-name()="oMathPara" or (local-name()="oMath" and not(ancestor::*[local-name()="oMathPara"]))]'
    ):
        text = _compact_formula_text(_omml_to_text(formula))
        if text:
            formulas.append(f"[FORMULA_OMML: {text}]")
    return formulas


def _extract_formula_image_markers(
    paragraph: Paragraph,
    formula_image_index: int,
) -> tuple[list[str], list[str], int, list[FormulaImageAsset]]:
    markers: list[str] = []
    formula_image_ids: list[str] = []
    assets: list[FormulaImageAsset] = []
    current_index = formula_image_index
    for image_index, relationship_id in enumerate(paragraph._p.xpath(".//a:blip/@r:embed"), start=1):
        image_part = paragraph.part.related_parts.get(relationship_id)
        filename = Path(str(image_part.partname)).name if image_part is not None else f"image-{image_index}"
        current_index += 1
        asset_id = f"formula-image-{current_index}"
        formula_image_ids.append(asset_id)
        markers.append(f"[FORMULA_IMAGE: {asset_id}; {filename}]")
        if image_part is not None:
            assets.append(FormulaImageAsset(asset_id=asset_id, filename=filename, content=image_part.blob))
    return markers, formula_image_ids, current_index, assets


def _omml_to_text(element) -> str:
    if element is None:
        return ""

    name = _local_name(element)

    if name == "t":
        return element.text or ""

    if name == "f":
        numerator = _first_child_by_name(element, "num")
        denominator = _first_child_by_name(element, "den")
        return f"({_omml_to_text(numerator)}) / ({_omml_to_text(denominator)})"

    if name == "sSup":
        base = _first_child_by_name(element, "e")
        superscript = _first_child_by_name(element, "sup")
        return f"{_omml_to_text(base)}^{_omml_to_text(superscript)}"

    if name == "sSub":
        base = _first_child_by_name(element, "e")
        subscript = _first_child_by_name(element, "sub")
        return f"{_omml_to_text(base)}_{_omml_to_text(subscript)}"

    if name == "sSubSup":
        base = _first_child_by_name(element, "e")
        subscript = _first_child_by_name(element, "sub")
        superscript = _first_child_by_name(element, "sup")
        return f"{_omml_to_text(base)}_{_omml_to_text(subscript)}^{_omml_to_text(superscript)}"

    if name in {"num", "den", "e", "lim", "sup", "sub"}:
        return "".join(_omml_to_text(child) for child in element)

    if name == "r":
        return "".join(_omml_to_text(child) for child in element)

    return " ".join(_omml_to_text(child) for child in element)


def _first_child_by_name(element, name: str):
    for child in element:
        if _local_name(child) == name:
            return child
    return None


def _local_name(element) -> str:
    if element is None:
        return ""
    return element.tag.rsplit("}", 1)[-1]


def _compact_formula_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _detect_heading_level(style_name: str, text: str, paragraph_index: int) -> int | None:
    if style_name.startswith("Heading"):
        return _parse_heading_level(style_name)

    if re.match(r"^Глава\s+\d+[\.\s]", text, flags=re.IGNORECASE):
        return 1

    if re.match(r"^Приложение\s+\d+", text, flags=re.IGNORECASE):
        return 1

    if style_name == "ConsPlusTitle" and paragraph_index > 8 and _looks_like_title(text):
        return 2

    return None


def _parse_heading_level(style_name: str) -> int | None:
    tokens = style_name.split()
    if len(tokens) < 2:
        return 1
    try:
        return int(tokens[-1])
    except ValueError:
        return 1


def _looks_like_title(text: str) -> bool:
    letters = [char for char in text if char.isalpha()]
    if not letters:
        return False

    upper_letters = [char for char in letters if char.isupper()]
    upper_ratio = len(upper_letters) / len(letters)
    return upper_ratio > 0.8 and len(text) <= 180
